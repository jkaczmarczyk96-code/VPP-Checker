from pathlib import Path
from io import BytesIO
from app.services.chunking import chunk_text
import os
import re

from fastapi import (
    APIRouter,
    UploadFile,
    File,
    Form,
    HTTPException,
    Cookie,
)

from docx import Document
from docx.shared import Pt
from pypdf import PdfReader

from app.supabase_client import supabase
from app.services.search import set_document_text
from app.services.storage import (
    add_document,
    get_documents,
    _write_all,
)

router = APIRouter(
    prefix="/api/v1/upload",
    tags=["Upload"]
)

COOKIE_NAME = "vpp_admin_session"
BUCKET = "documents"


# =========================
# AUTH
# =========================
def require_admin(
    session: str = Cookie(
        default=None,
        alias=COOKIE_NAME
    )
):
    if session != "logged_in":
        raise HTTPException(
            status_code=401,
            detail="Unauthorized"
        )


# =========================
# HELPERS
# =========================
def sanitize_filename(filename: str) -> str:
    filename = filename.strip()
    filename = filename.replace(" ", "_")
    filename = re.sub(
        r"[^a-zA-Z0-9._-]",
        "",
        filename
    )
    return filename


def format_paragraph(
    paragraph,
    size,
    bold=False,
    italic=False
):
    for run in paragraph.runs:
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic


def looks_like_part(text: str) -> bool:
    patterns = [
        r"^(Část|Časť|Oddíl|Oddiel|Sekce|Kapitola|Hlava)\b",
        r"^\d+\.\s+[A-ZÁČĎÉĚÍĹĽŇÓÔŘŠŤÚŮÝŽ]",
    ]
    return any(
        re.search(p, text, re.IGNORECASE)
        for p in patterns
    )


def looks_like_article(text: str) -> bool:
    return bool(
        re.search(
            r"^(Článek|Článok|Čl\.)\s*\d+",
            text,
            re.IGNORECASE
        )
    )


def looks_like_short_heading(text: str) -> bool:
    if len(text) > 80:
        return False

    if text.endswith(";"):
        return False

    return len(text.split()) <= 8


# =========================
# DOCX NORMALIZER
# =========================
def normalize_docx_bytes(content: bytes) -> bytes:
    doc = Document(BytesIO(content))

    for p in doc.paragraphs:
        text = p.text.strip()

        if not text:
            continue

        if looks_like_part(text):
            format_paragraph(
                p, 14, bold=True
            )

        elif looks_like_article(text):
            format_paragraph(
                p, 13, bold=True
            )

        elif looks_like_short_heading(text):
            format_paragraph(
                p, 12, bold=True
            )

        else:
            format_paragraph(
                p, 11
            )

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()


def parse_docx(content: bytes) -> str:
    normalized = normalize_docx_bytes(content)

    doc = Document(BytesIO(normalized))

    lines = []

    for p in doc.paragraphs:
        text = p.text.strip()

        if text:
            lines.append(text)

    return "\n".join(lines)


def parse_pdf(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))

    pages = []

    for i, page in enumerate(
        reader.pages,
        start=1
    ):
        page_text = page.extract_text()

        if page_text:
            pages.append(
                f"[[PAGE:{i}]]\n{page_text.strip()}"
            )

    return "\n\n".join(pages)


def remove_old_record(filename: str):
    docs = get_documents()

    docs = [
        d for d in docs
        if d.get("file_name") != filename
    ]

    _write_all(docs)

    try:
        supabase.storage.from_(BUCKET).remove(
            [filename]
        )
    except Exception:
        pass


# =========================
# LIST FILES
# =========================
@router.get("/")
async def list_files():
    docs = get_documents()

    return {
        "files": sorted([
            d.get("file_name")
            for d in docs
            if d.get("file_name")
        ])
    }


# =========================
# UPLOAD
# =========================
@router.post("/")
async def upload_file(
    insurer: str = Form(...),
    document_title: str = Form(...),
    file: UploadFile = File(...),
    session: str = Cookie(
        default=None,
        alias=COOKIE_NAME
    ),
):
    require_admin(session)

    if not insurer.strip():
        raise HTTPException(
            status_code=400,
            detail="Pojišťovna je povinná."
        )

    if not document_title.strip():
        raise HTTPException(
            status_code=400,
            detail="Název dokumentu je povinný."
        )

    safe_name = sanitize_filename(
        file.filename
    )

    if not safe_name:
        raise HTTPException(
            status_code=400,
            detail="Neplatný název souboru."
        )

    filename_lower = safe_name.lower()

    if not (
        filename_lower.endswith(".pdf")
        or filename_lower.endswith(".docx")
    ):
        raise HTTPException(
            status_code=400,
            detail="Podporován je pouze DOCX nebo PDF."
        )

    content = await file.read()

    remove_old_record(safe_name)

    # =====================
    # Upload do Supabase Storage
    # =====================
    try:
        supabase.storage.from_(BUCKET).upload(
            path=safe_name,
            file=content,
            file_options={
                "content-type": file.content_type,
                "upsert": "true"
            }
        )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Upload failed: {str(e)}"
        )

    # =====================
    # Parse
    # =====================
    if filename_lower.endswith(".docx"):
        text = parse_docx(content)
    else:
        text = parse_pdf(content)

    if not text.strip():
        raise HTTPException(
            status_code=400,
            detail="Nepodařilo se načíst text dokumentu."
        )

    # =====================
    # Chunks
    # =====================
    set_document_text(text)

    chunks = chunk_text(text)

    add_document({
        "insurer": insurer.strip(),
        "document_title": document_title.strip(),
        "file_name": safe_name,
        "text_content": text,
        "chunks": chunks
    })

    return {
        "filename": safe_name,
        "insurer": insurer,
        "document_title": document_title,
        "status": "uploaded",
        "chunks": len(chunks),
        "text_preview": text[:300]
    }


# =========================
# DELETE
# =========================
@router.delete("/{filename}")
async def delete_file(
    filename: str,
    session: str = Cookie(
        default=None,
        alias=COOKIE_NAME
    ),
):
    require_admin(session)

    safe_name = sanitize_filename(
        filename
    )

    remove_old_record(safe_name)

    return {
        "status": "deleted",
        "filename": safe_name
    }
